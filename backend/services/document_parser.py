import docx
import PyPDF2
import pdfplumber
from typing import Optional
import io

class DocumentParser:
    """Parse DOCX and PDF documents to extract text"""
    
    MAX_TEXT_LENGTH = 100000  # Maximum text length to prevent memory issues
    
    @staticmethod
    def parse_docx(file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(io.BytesIO(file_content))
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())
            
            full_text = "\n\n".join(text_parts)
            # Limit text length to prevent memory issues
            if len(full_text) > DocumentParser.MAX_TEXT_LENGTH:
                full_text = full_text[:DocumentParser.MAX_TEXT_LENGTH]
                print(f"Warning: DOCX text truncated to {DocumentParser.MAX_TEXT_LENGTH} characters")
            return full_text
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {str(e)}")
    
    @staticmethod
    def parse_pdf(file_content: bytes) -> str:
        """Extract text from PDF file using pdfplumber (better for complex layouts)"""
        try:
            text_parts = []
            
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text and text.strip():
                        text_parts.append(text.strip())
            
            if not text_parts:
                # Fallback to PyPDF2 if pdfplumber fails
                full_text = DocumentParser._parse_pdf_pypdf2(file_content)
            else:
                full_text = "\n\n".join(text_parts)
            
            # Limit text length to prevent memory issues
            if len(full_text) > DocumentParser.MAX_TEXT_LENGTH:
                full_text = full_text[:DocumentParser.MAX_TEXT_LENGTH]
                print(f"Warning: PDF text truncated to {DocumentParser.MAX_TEXT_LENGTH} characters")
            return full_text
        except Exception as e:
            # Try fallback method
            try:
                return DocumentParser._parse_pdf_pypdf2(file_content)
            except:
                raise ValueError(f"Failed to parse PDF: {str(e)}")
    
    @staticmethod
    def _parse_pdf_pypdf2(file_content: bytes) -> str:
        """Fallback PDF parser using PyPDF2"""
        text_parts = []
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text and text.strip():
                text_parts.append(text.strip())
        
        full_text = "\n\n".join(text_parts)
        # Limit text length to prevent memory issues
        if len(full_text) > DocumentParser.MAX_TEXT_LENGTH:
            full_text = full_text[:DocumentParser.MAX_TEXT_LENGTH]
            print(f"Warning: PDF (PyPDF2) text truncated to {DocumentParser.MAX_TEXT_LENGTH} characters")
        return full_text
    
    @staticmethod
    def parse_document(file_content: bytes, filename: str) -> str:
        """Parse document based on file extension"""
        filename_lower = filename.lower()
        
        if filename_lower.endswith('.docx'):
            return DocumentParser.parse_docx(file_content)
        elif filename_lower.endswith('.pdf'):
            return DocumentParser.parse_pdf(file_content)
        else:
            raise ValueError(f"Unsupported file type. Only .docx and .pdf are supported.")
    
    @staticmethod
    def validate_file(filename: str, file_size: int, max_size_mb: int = 10) -> tuple[bool, Optional[str]]:
        """Validate file type and size"""
        # Check file extension
        filename_lower = filename.lower()
        if not (filename_lower.endswith('.docx') or filename_lower.endswith('.pdf')):
            return False, "Only .docx and .pdf files are supported"
        
        # Check file size
        max_size_bytes = max_size_mb * 1024 * 1024
        if file_size > max_size_bytes:
            return False, f"File size exceeds {max_size_mb}MB limit"
        
        return True, None
